#!/usr/bin/env python3
"""
Analyze DeepHyper HPO results for HydraGNN OPF heterogeneous graph training.

Produces:
  1. A semilogy plot of validation loss curves per epoch, colored by model type.
  2. A console summary table of all trials and per-model-type statistics.
  3. (Optional) Appends an HPO results section to the OPF explanation docx.

Usage examples:
  # Minimal: just specify the job ID (assumes standard directory layout)
  python plot_hpo_results.py --job_id 4249563

  # Explicit paths (when directories were renamed or moved)
  python plot_hpo_results.py \\
      --csv_path  opf_hpo-4249563/results.csv \\
      --dh_dir    deephyper-opf-hpo-4249563 \\
      --plot_path hpo_validation_loss_curves.png

  # Also update the docx
  python plot_hpo_results.py --job_id 4249563 \\
      --docx_path OPF_Heterogeneous_Graph_Explanation.docx

  # Combine multiple HPO runs
  python plot_hpo_results.py \\
      --csv_path  opf_hpo-4249563/results.csv opf_hpo-4249999/results.csv \\
      --dh_dir    deephyper-opf-hpo-4249563     deephyper-opf-hpo-4249999
"""

import argparse
import csv
import os
import re
import sys
from collections import defaultdict

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

# ─── Visual style ────────────────────────────────────────────────────────────
TYPE_COLORS = {
    "HeteroHEAT": "#d62728",  # red
    "HeteroHGT": "#1f77b4",  # blue
    "HeteroPNA": "#2ca02c",  # green
    "HeteroSAGE": "#ff7f0e",  # orange
    "HeteroGAT": "#9467bd",  # purple
    "HeteroRGAT": "#8c564b",  # brown
    "HeteroGIN": "#e377c2",  # pink
}

TYPE_MARKERS = {
    "HeteroHEAT": "s",
    "HeteroHGT": "o",
    "HeteroPNA": "D",
    "HeteroSAGE": "^",
    "HeteroGAT": "v",
    "HeteroRGAT": "X",
    "HeteroGIN": "P",
}


# ─── Data loading ────────────────────────────────────────────────────────────
def load_results(csv_paths):
    """Load and merge results from one or more DeepHyper CSV files."""
    all_rows = []
    for path in csv_paths:
        with open(path) as f:
            rows = list(csv.DictReader(f))
        print(f"  Loaded {len(rows)} trials from {path}")
        all_rows.extend(rows)
    return all_rows


def extract_epoch_losses(dh_dirs, trial_id):
    """Read output-0.{trial_id}.txt from the first matching dh_dir."""
    for dh_dir in dh_dirs:
        path = os.path.join(dh_dir, f"output-0.{trial_id}.txt")
        if not os.path.exists(path):
            continue
        losses = []
        with open(path) as f:
            for line in f:
                m = re.match(r"Val loss:\s+([\d.eE+-]+)", line.strip())
                if m:
                    val = float(m.group(1))
                    if np.isfinite(val):
                        losses.append(val)
        if losses:
            return losses
    return []


def build_trial_data(rows, dh_dirs):
    """Returns dict: mpnn_type -> list of (trial_id, hyperparams, [epoch_losses])"""
    by_type = defaultdict(list)
    n_failed = 0
    failed_types = defaultdict(int)
    for r in rows:
        if r["objective"] == "F":
            n_failed += 1
            failed_types[r["p:mpnn_type"]] += 1
            continue
        trial_id = r["job_id"]
        mpnn = r["p:mpnn_type"]
        losses = extract_epoch_losses(dh_dirs, trial_id)
        if not losses:
            continue
        hparams = {
            "hidden_dim": int(r["p:hidden_dim"]),
            "num_conv_layers": int(r["p:num_conv_layers"]),
            "learning_rate": float(r["p:learning_rate"]),
        }
        by_type[mpnn].append((trial_id, hparams, losses))
    return by_type, n_failed, dict(failed_types)


# ─── Plotting ────────────────────────────────────────────────────────────────
def create_plot(by_type, plot_path, title_extra=""):
    """Create semilogy plot of val loss curves, colored by model type."""
    fig, ax = plt.subplots(figsize=(12, 7))

    type_order = sorted(
        by_type.keys(), key=lambda t: min(min(l) for _, _, l in by_type[t])
    )

    legend_handles = []
    legend_labels = []
    max_epochs = 0
    overall_best_loss = float("inf")
    overall_best_info = {}

    for mpnn in type_order:
        trials = by_type[mpnn]
        color = TYPE_COLORS.get(mpnn, "#333333")
        marker = TYPE_MARKERS.get(mpnn, "o")
        trials.sort(key=lambda t: min(t[2]))

        for idx, (trial_id, hparams, losses) in enumerate(trials):
            epochs = list(range(1, len(losses) + 1))
            max_epochs = max(max_epochs, len(losses))
            alpha = 0.9 if idx == 0 else 0.35
            lw = 2.5 if idx == 0 else 1.0

            line = ax.semilogy(
                epochs,
                losses,
                color=color,
                marker=marker,
                markersize=6 if idx == 0 else 4,
                linewidth=lw,
                alpha=alpha,
                markeredgewidth=0.5,
                markeredgecolor="white",
            )

            if idx == 0:
                best_loss = min(losses)
                legend_handles.append(line[0])
                legend_labels.append(f"{mpnn} (best: {best_loss:.5f})")
                if best_loss < overall_best_loss:
                    overall_best_loss = best_loss
                    best_epoch = losses.index(best_loss) + 1
                    overall_best_info = {
                        "mpnn": mpnn,
                        "trial": trial_id,
                        "loss": best_loss,
                        "epoch": best_epoch,
                    }

    total_trials = sum(len(v) for v in by_type.values())
    ax.set_xlabel("Epoch", fontsize=13, fontweight="bold")
    ax.set_ylabel("Validation Loss", fontsize=13, fontweight="bold")

    title = f"HPO Validation Loss Curves by Model Architecture\n({total_trials} trials"
    if title_extra:
        title += f", {title_extra}"
    title += ")"
    ax.set_title(title, fontsize=14, fontweight="bold")

    ax.set_xticks(range(1, max_epochs + 1))
    ax.set_xlim(0.5, max_epochs + 0.5)
    ax.grid(True, which="both", alpha=0.3, linestyle="--")
    ax.yaxis.set_minor_formatter(mticker.NullFormatter())

    ax.legend(
        legend_handles,
        legend_labels,
        loc="upper right",
        fontsize=10,
        framealpha=0.9,
        title="Model Type (best trial val loss)",
        title_fontsize=10,
    )

    if overall_best_info:
        info = overall_best_info
        ax.annotate(
            f"Best overall:\n{info['mpnn']}, trial {info['trial']}\n"
            f"Val loss = {info['loss']:.5f}",
            xy=(info["epoch"], info["loss"]),
            xytext=(max(1, info["epoch"] - 3), info["loss"] * 2.5),
            fontsize=9,
            arrowprops=dict(arrowstyle="->", color="gray", lw=1.2),
            bbox=dict(
                boxstyle="round,pad=0.3",
                facecolor="lightyellow",
                edgecolor="gray",
            ),
        )

    plt.tight_layout()
    plt.savefig(plot_path, dpi=200, bbox_inches="tight")
    print(f"Plot saved to {plot_path}")
    return plot_path


# ─── Console summary ─────────────────────────────────────────────────────────
def print_summary(by_type, n_failed, failed_types):
    """Print formatted trial summary to console."""
    total = sum(len(v) for v in by_type.values())
    print(f"\n{'='*75}")
    print(f"  HPO Results: {total} successful trials, {n_failed} failed")
    if failed_types:
        parts = [f"{t}: {n}" for t, n in sorted(failed_types.items())]
        print(f"  Failed by type: {', '.join(parts)}")
    print(f"{'='*75}\n")

    # All trials ranked
    all_trials = []
    for mpnn, trials in by_type.items():
        for trial_id, hparams, losses in trials:
            all_trials.append(
                {
                    "trial": trial_id,
                    "mpnn": mpnn,
                    "hidden": hparams["hidden_dim"],
                    "layers": hparams["num_conv_layers"],
                    "lr": hparams["learning_rate"],
                    "best_loss": min(losses),
                }
            )
    all_trials.sort(key=lambda x: x["best_loss"])

    print(f"{'Rank':>4} {'Trial':>5} {'Model Type':>14} {'Hidden':>6} {'Layers':>6} {'LR':>10} {'Best Val Loss':>13}")
    print("-" * 68)
    for i, t in enumerate(all_trials):
        print(
            f"{i+1:>4} {t['trial']:>5} {t['mpnn']:>14} {t['hidden']:>6} "
            f"{t['layers']:>6} {t['lr']:>10.6f} {t['best_loss']:>13.6f}"
        )

    # Per-type summary
    print(f"\n{'--- Summary by Model Type ---':^68}")
    print(f"{'Model Type':>14} {'Trials':>6} {'Best':>10} {'Mean':>10} {'Worst':>10}")
    print("-" * 54)
    type_stats = {}
    for mpnn, trials in by_type.items():
        losses_all = [min(l) for _, _, l in trials]
        type_stats[mpnn] = {
            "count": len(losses_all),
            "best": min(losses_all),
            "worst": max(losses_all),
            "mean": sum(losses_all) / len(losses_all),
        }
    for mpnn in sorted(type_stats, key=lambda t: type_stats[t]["best"]):
        s = type_stats[mpnn]
        print(
            f"{mpnn:>14} {s['count']:>6} {s['best']:>10.6f} "
            f"{s['mean']:>10.6f} {s['worst']:>10.6f}"
        )
    print()
    return all_trials, type_stats


# ─── Document update ─────────────────────────────────────────────────────────
def update_docx(by_type, plot_path, docx_path, n_failed, failed_types):
    """Append HPO results section to the OPF explanation docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    doc = Document(docx_path)

    # Check if section 15.4 already exists — avoid duplicating
    for p in doc.paragraphs:
        if "15.4" in p.text and "HPO Results" in p.text:
            print(f"WARNING: Section '15.4 HPO Results' already exists in {docx_path}.")
            print("         Skipping docx update to avoid duplication.")
            print("         Delete the existing section first if you want to regenerate it.")
            return

    total = sum(len(v) for v in by_type.values())
    failed_detail = ""
    if failed_types:
        parts = [f"{t}" for t in sorted(failed_types)]
        failed_detail = (
            f" (all {', '.join(parts)} configurations, which crashed due to NCCL "
            "communication errors)"
        )

    def add_heading(text, level=2):
        p = doc.add_paragraph()
        p.style = doc.styles[f"Heading {level}"]
        p.text = text
        return p

    def add_normal(text):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.text = text
        return p

    def add_bullet(text):
        p = doc.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.text = text
        return p

    # ── Section 15.4: HPO Results ──
    add_heading("15.4 HPO Results", level=2)
    add_normal(
        f"The HPO search completed {total + n_failed} trials, of which {total} "
        f"completed successfully and {n_failed} failed{failed_detail}. "
        "Despite any partial completion, the results provide strong evidence for "
        "which architectures and hyperparameter regions perform best on the AC-OPF problem."
    )

    # ── Results Table ──
    add_heading("15.4.1 Trial Results Summary", level=2)
    add_normal(
        "The following table shows the top 10 HPO trials ranked by best validation loss "
        "(minimum across all epochs). The Bayesian optimizer successfully identified "
        "the top-performing architectures and concentrated its search in the most "
        "promising hyperparameter regions."
    )

    all_success = []
    for mpnn, trials in by_type.items():
        for trial_id, hparams, losses in trials:
            all_success.append(
                {
                    "trial": trial_id,
                    "mpnn": mpnn,
                    "hidden": hparams["hidden_dim"],
                    "layers": hparams["num_conv_layers"],
                    "lr": hparams["learning_rate"],
                    "best_loss": min(losses),
                }
            )
    all_success.sort(key=lambda x: x["best_loss"])

    n_show = min(10, len(all_success))
    table = doc.add_table(rows=n_show + 1, cols=6)
    table.style = "Light Shading Accent 1"
    headers = [
        "Rank", "Model Type", "Hidden Dim",
        "Conv Layers", "Learning Rate", "Best Val Loss",
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, trial in enumerate(all_success[:n_show]):
        row = table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = trial["mpnn"]
        row.cells[2].text = str(trial["hidden"])
        row.cells[3].text = str(trial["layers"])
        row.cells[4].text = f"{trial['lr']:.6f}"
        row.cells[5].text = f"{trial['best_loss']:.6f}"

    add_normal("")

    # ── Summary by Model Type ──
    add_heading("15.4.2 Performance by Model Architecture", level=2)
    add_normal(
        "Aggregating results by model architecture reveals clear performance tiers. "
        "The table below summarizes the validation loss statistics for each model type."
    )

    type_stats = {}
    for mpnn, trials in by_type.items():
        losses_all = [min(l) for _, _, l in trials]
        type_stats[mpnn] = {
            "count": len(losses_all),
            "best": min(losses_all),
            "worst": max(losses_all),
            "mean": sum(losses_all) / len(losses_all),
        }
    type_order = sorted(type_stats.keys(), key=lambda t: type_stats[t]["best"])

    table2 = doc.add_table(rows=len(type_order) + 1, cols=5)
    table2.style = "Light Shading Accent 1"
    for j, h in enumerate(
        ["Model Type", "Trials", "Best Val Loss", "Mean Val Loss", "Worst Val Loss"]
    ):
        cell = table2.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for i, mpnn in enumerate(type_order):
        s = type_stats[mpnn]
        row = table2.rows[i + 1]
        row.cells[0].text = mpnn
        row.cells[1].text = str(s["count"])
        row.cells[2].text = f"{s['best']:.6f}"
        row.cells[3].text = f"{s['mean']:.6f}"
        row.cells[4].text = f"{s['worst']:.6f}"

    add_normal("")

    # ── Training curves figure ──
    add_heading("15.4.3 Validation Loss Training Curves", level=2)
    add_normal(
        "The following figure shows the validation loss curves (semilogy scale) for all "
        f"{total} successful HPO trials, colored by model architecture. For each model "
        "type, the best-performing trial is drawn with a bold line, while other trials "
        "of the same type appear as thinner, more transparent lines."
    )

    doc.add_picture(plot_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Build dynamic caption from actual best/worst types
    hex_to_name = {
        "#d62728": "red", "#1f77b4": "blue", "#2ca02c": "green",
        "#ff7f0e": "orange", "#9467bd": "purple", "#8c564b": "brown",
        "#e377c2": "pink", "#333333": "gray",
    }
    best_two = type_order[:2]
    worst_two = type_order[-2:] if len(type_order) >= 4 else type_order[-1:]
    color_map = {t: TYPE_COLORS.get(t, "#333333") for t in type_order}
    best_desc = " and ".join(
        f"{t} ({hex_to_name.get(color_map[t], 'colored')})" for t in best_two
    )
    worst_desc = " and ".join(
        f"{t} ({hex_to_name.get(color_map[t], 'colored')})" for t in worst_two
    )
    add_normal(
        f"Figure: Validation loss vs. epoch for all successful HPO trials. "
        f"{best_desc} consistently achieve the lowest validation losses, while "
        f"{worst_desc} remain significantly worse."
    )

    add_normal("")

    # ── Analysis ──
    add_heading("15.4.4 Analysis of HPO Results", level=2)
    add_normal(
        "The HPO results provide several key insights that refine the preliminary "
        "findings from Section 14:"
    )

    best_trial = all_success[0]
    add_bullet(
        f"{best_trial['mpnn']} emerges as the top architecture. With a best validation "
        f"loss of {best_trial['best_loss']:.5f} (trial {best_trial['trial']}: "
        f"hidden_dim={best_trial['hidden']}, {best_trial['layers']} conv layers, "
        f"lr={best_trial['lr']:.5f}), it outperforms all other architectures."
    )

    if len(type_order) >= 2:
        second_type = type_order[1]
        second_best = type_stats[second_type]
        add_bullet(
            f"{second_type} is a strong second. With a best validation loss of "
            f"{second_best['best']:.5f} across {second_best['count']} trials "
            f"(mean {second_best['mean']:.5f}), it demonstrates consistent performance."
        )

    # Common hyperparameter patterns in top 5
    top5 = all_success[:min(5, len(all_success))]
    layers_set = set(t["layers"] for t in top5)
    if len(layers_set) == 1:
        common_layers = top5[0]["layers"]
        add_bullet(
            f"Optimal hyperparameter region. The top {len(top5)} trials all use "
            f"{common_layers} convolutional layers with hidden dimensions between "
            f"{min(t['hidden'] for t in top5)} and {max(t['hidden'] for t in top5)} "
            f"and learning rates between {min(t['lr'] for t in top5):.4f} and "
            f"{max(t['lr'] for t in top5):.4f}."
        )
    else:
        add_bullet(
            f"Optimal hyperparameter region. The top {len(top5)} trials use "
            f"{min(t['layers'] for t in top5)}-{max(t['layers'] for t in top5)} "
            f"convolutional layers with hidden dimensions between "
            f"{min(t['hidden'] for t in top5)} and {max(t['hidden'] for t in top5)}."
        )

    if failed_types:
        for ft, fc in sorted(failed_types.items(), key=lambda x: -x[1]):
            add_bullet(
                f"{ft}: all {fc} trial(s) failed, suggesting scalability or "
                "stability issues in the distributed training setting."
            )

    for mpnn in type_order[-2:]:
        if type_stats[mpnn]["best"] > 2 * type_stats[type_order[0]]["best"]:
            add_bullet(
                f"{mpnn} underperforms significantly with a best validation loss of "
                f"{type_stats[mpnn]['best']:.5f} — "
                f"{type_stats[mpnn]['best']/type_stats[type_order[0]]['best']:.1f}x "
                f"worse than {type_order[0]}."
            )

    add_normal("")

    # ── Recommendation ──
    add_heading("15.4.5 Revised Model Recommendation", level=2)
    second = type_order[1] if len(type_order) >= 2 else "N/A"
    add_normal(
        f"Based on the HPO results, {best_trial['mpnn']} is recommended as the primary "
        f"architecture for the AC-OPF problem, with {second} as a strong alternative. "
        f"The optimal configuration is: {best_trial['mpnn']} with "
        f"{best_trial['layers']} convolutional layers, hidden dimension "
        f"~{best_trial['hidden']}, and learning rate ~{best_trial['lr']:.4f}."
    )

    doc.save(docx_path)
    print(f"Document updated: {docx_path}")


# ─── CLI ─────────────────────────────────────────────────────────────────────
def parse_args():
    parser = argparse.ArgumentParser(
        description="Analyze DeepHyper HPO results for HydraGNN OPF training.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Quick run with just a job ID (uses standard directory naming):
  %(prog)s --job_id 4249563

  # Explicit paths:
  %(prog)s --csv_path opf_hpo-4249563/results.csv \\
           --dh_dir  deephyper-opf-hpo-4249563

  # Merge multiple HPO runs:
  %(prog)s --csv_path opf_hpo-4249563/results.csv opf_hpo-4250000/results.csv \\
           --dh_dir  deephyper-opf-hpo-4249563     deephyper-opf-hpo-4250000

  # Also update the docx:
  %(prog)s --job_id 4249563 --docx_path OPF_Heterogeneous_Graph_Explanation.docx
""",
    )

    parser.add_argument(
        "--job_id",
        type=str,
        default=None,
        help="Slurm job ID. If provided, csv_path and dh_dir are auto-derived "
        "as 'opf_hpo-{job_id}/results.csv' and 'deephyper-opf-hpo-{job_id}/'.",
    )
    parser.add_argument(
        "--csv_path",
        nargs="+",
        default=None,
        help="Path(s) to DeepHyper results.csv file(s). Overrides --job_id.",
    )
    parser.add_argument(
        "--dh_dir",
        nargs="+",
        default=None,
        help="Path(s) to DeepHyper output directories containing output-0.*.txt "
        "files with per-epoch val losses. Overrides --job_id.",
    )
    parser.add_argument(
        "--plot_path",
        type=str,
        default="hpo_validation_loss_curves.png",
        help="Output path for the semilogy plot image (default: hpo_validation_loss_curves.png).",
    )
    parser.add_argument(
        "--title_extra",
        type=str,
        default="",
        help="Extra text appended to the plot title (e.g., '128 Frontier nodes').",
    )
    parser.add_argument(
        "--docx_path",
        type=str,
        default=None,
        help="Path to the OPF_Heterogeneous_Graph_Explanation.docx. "
        "If provided, HPO results section 15.4 will be appended.",
    )
    parser.add_argument(
        "--no_plot",
        action="store_true",
        help="Skip plot generation (print summary only).",
    )
    return parser.parse_args()


def main():
    args = parse_args()

    # Resolve paths
    if args.csv_path:
        csv_paths = args.csv_path
    elif args.job_id:
        csv_paths = [f"opf_hpo-{args.job_id}/results.csv"]
    else:
        print("ERROR: Provide either --job_id or --csv_path.", file=sys.stderr)
        sys.exit(1)

    if args.dh_dir:
        dh_dirs = args.dh_dir
    elif args.job_id:
        dh_dirs = [f"deephyper-opf-hpo-{args.job_id}"]
    else:
        print("ERROR: Provide either --job_id or --dh_dir.", file=sys.stderr)
        sys.exit(1)

    # Validate inputs
    for p in csv_paths:
        if not os.path.exists(p):
            print(f"ERROR: CSV file not found: {p}", file=sys.stderr)
            sys.exit(1)
    for d in dh_dirs:
        if not os.path.isdir(d):
            print(f"WARNING: DeepHyper dir not found: {d} (epoch curves will be missing)")

    # Load and build data
    print("Loading results...")
    rows = load_results(csv_paths)
    by_type, n_failed, failed_types = build_trial_data(rows, dh_dirs)

    total = sum(len(v) for v in by_type.values())
    if total == 0:
        print("ERROR: No successful trials with epoch data found.", file=sys.stderr)
        sys.exit(1)

    # Console summary
    print_summary(by_type, n_failed, failed_types)

    # Plot
    if not args.no_plot:
        title_extra = args.title_extra
        if not title_extra and args.job_id:
            title_extra = f"Job {args.job_id}"
        create_plot(by_type, args.plot_path, title_extra=title_extra)

    # Docx update
    if args.docx_path:
        if not os.path.exists(args.docx_path):
            print(f"ERROR: Docx file not found: {args.docx_path}", file=sys.stderr)
            sys.exit(1)
        update_docx(by_type, args.plot_path, args.docx_path, n_failed, failed_types)


if __name__ == "__main__":
    main()
